import os
import docx
import time
import numpy
from PIL import Image
from datetime import datetime

exec_Time=datetime.now().strftime("%r").replace(":",";")
exec_Date=datetime.now().strftime("%d%h%G")
doc=docx.Document()
sec=doc.sections
for x in sec:
    x.top_margin=docx.shared.Cm(1.27)
    x.bottom_margin=docx.shared.Cm(1.27)
    x.right_margin=docx.shared.Cm(1.27)
    x.left_margin=docx.shared.Cm(1.27)

s_=' '
gs="@$#&%*+=:^;~-.{}".format(s_)
os.chdir("/Images")
for i in os.walk(os.getcwd()):
    dir_=i[0]
    images=i[2]
imgc=0
pages=[]
T_data=[]
for pic in images:
    imgc+=1
    print("\n\n\n-->",pic)
    loc=dir_+"\\"+pic
    img=Image.open(loc).convert('L')
    ar=numpy.array(img)
    pic_w,pic_h=ar.shape
    avg_data=numpy.average(ar.reshape(pic_w,pic_h))
    types=[80,150,250,pic_w]
    imgSet=[]
    T_type=0
    cols_w=0
    Doc_t=0
    for cols in types:
        T_Col1=time.time()
        if cols>pic_w:
            cols=pic_w
            cols_w+=1
            types.pop()
        if cols_w<=1:
            scale=0.47
            width=pic_w/cols
            height=width/scale
            rows=int(pic_h/height)
            imgList=[]
            for i in range(rows):
                y=int(i*height)
                y_=int((i+1)*height)
                imgList.append("")
                Val=""
                for j in range(cols):
                    x=int(j*width)
                    x_=int((j+1)*width)
                    img_=img.crop((x,y,x_,y_))
                    ar=numpy.array(img_)
                    w,h=ar.shape
                    avg_data=numpy.average(ar.reshape(w,h))
                    gsVal=gs[int((avg_data*14)/255)]
                    Val+=gsVal
                Val='\n'+Val
                if len(set(list(Val)))==2:
                    pass
                else:
                    imgList[i]=Val
            imgSet+=imgList+['<br><br><br><br><br>']
            imgData=''.join(imgList)
            size_dict={50:7.5,100:5.5,150:4.5,200:3.5,250:3,\
            300:2.5,350:2,400:1.5,450:1} 
            for y in size_dict.keys():
                if cols/y<=1:
                    px=size_dict[y]
                    break
                else:
                    px=size_dict[450]
            T_Col2=time.time()
            chrC=rows*cols
            T_type=T_Col2-T_Col1
            print("Resolution Type: ",cols," Done!")
            print(">No. Of Characters: ", chrC)
            print("Execution time: ",T_type,"Secs\n")
            T_doc1=time.time()
            doc.add_heading("--->"+pic+" CharacterCount: "+str(chrC))
            doc.add_paragraph("\n\n\n\n\n\n\n\n")
            para=doc.add_paragraph()
            p=para.add_run(imgData)
            p.font.size=docx.shared.Pt(px)
            p.font.name='Courier New'
            para.paragraph_format.line_spacing=docx.shared.Cm(0)
            doc.add_page_break()
            T_doc2=time.time()
            T_doc=(T_doc2-T_doc1)
            Doc_t+=T_doc
    print("\nWord Doc ",pic," Written!")
    print("Time Consumed: ",Doc_t,"Secs\n")
    
doc.save("DaWm.docx")